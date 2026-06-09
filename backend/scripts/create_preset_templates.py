"""生成 3 个预置签字页 Word 模板（含 {{key}} 占位符）。"""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

OUTPUT_DIR = Path(__file__).resolve().parent.parent / "app" / "templates"


def _add_centered(doc: Document, text: str, *, bold: bool = False) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.font.size = Pt(12)
    run.bold = bold


def create_law_firm_template(path: Path) -> None:
    doc = Document()
    _add_centered(
        doc,
        "（本页无正文，为《{{target_company_name}}向{{target_investor_type}}公开发行股票"
        "并在{{exchange_name}}上市之{{document_type}}》之签署页）",
    )
    doc.add_paragraph()
    _add_centered(doc, "{{law_firm_name}}", bold=True)
    doc.add_paragraph()
    _add_centered(doc, "律师事务所负责人：{{law_firm_director}}")
    _add_centered(doc, "经办律师：{{handling_lawyer}}")
    _add_centered(doc, "经办律师：{{handling_lawyer}}")
    doc.add_paragraph()
    _add_centered(doc, "{{signing_date}}")
    doc.save(path)


def create_natural_shareholder_template(path: Path) -> None:
    doc = Document()
    _add_centered(
        doc,
        "（本页无正文，为《{{target_company_name}}{{meeting_year}}年第{{meeting_session}}次"
        "临时股东大会会议决议》之签字页）",
    )
    doc.add_paragraph()
    _add_centered(doc, "{{natural_shareholder_name}}（自然人股东适用）")
    doc.add_paragraph()
    _add_centered(doc, "签字：")
    doc.add_paragraph()
    _add_centered(doc, "{{target_company_name}}")
    _add_centered(doc, "{{signing_date}}")
    doc.save(path)


def create_institutional_shareholder_template(path: Path) -> None:
    doc = Document()
    _add_centered(
        doc,
        "（本页无正文，为《{{target_company_name}}{{meeting_year}}年第{{meeting_session}}次"
        "临时股东大会会议决议》之签字页）",
    )
    doc.add_paragraph()
    _add_centered(doc, "{{institutional_shareholder_name}}（公章）")
    doc.add_paragraph()
    _add_centered(doc, "执行事务合伙人委派代表/授权代表（签字）：")
    _add_centered(doc, "姓名：{{authorized_representative_name}}")
    doc.add_paragraph()
    _add_centered(doc, "{{target_company_name}}")
    _add_centered(doc, "{{signing_date}}")
    doc.save(path)


def main() -> None:
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    create_law_firm_template(OUTPUT_DIR / "law_firm_signing_page.docx")
    create_natural_shareholder_template(OUTPUT_DIR / "natural_shareholder_signing_page.docx")
    create_institutional_shareholder_template(
        OUTPUT_DIR / "institutional_shareholder_signing_page.docx"
    )
    print(f"Created 3 preset templates in {OUTPUT_DIR}")


if __name__ == "__main__":
    main()
